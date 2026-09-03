"""Generate Ubuzima Bwiza project defense Word document."""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

NAVY = RGBColor(0x0F, 0x3D, 0x5E)
TEAL = RGBColor(0x0D, 0x6E, 0x6E)
DARK = RGBColor(0x1A, 0x1A, 0x1A)
GRAY = RGBColor(0x44, 0x44, 0x44)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "0F3D5E"
ALT_BG = "E8F2F4"


def set_run_font(run, name="Times New Roman", size=12, bold=False, italic=False, color=DARK):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color


def shade_cell(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    shd.set(qn("w:val"), "clear")
    tcPr.append(shd)


def set_cell_borders(cell):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "0F3D5E")
        tcBorders.append(el)
    tcPr.append(tcBorders)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fldChar1 = OxmlElement("w:fldChar")
    fldChar1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    fldChar2 = OxmlElement("w:fldChar")
    fldChar2.set(qn("w:fldCharType"), "end")
    run._r.append(fldChar1)
    run._r.append(instr)
    run._r.append(fldChar2)
    set_run_font(run, size=10, color=GRAY)


def prevent_table_split(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
    tblLook = OxmlElement("w:tblLook")
    tblLook.set(qn("w:firstRow"), "1")
    tblPr.append(tblLook)


class Report:
    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()
        self._setup_header_footer()

    def _setup_page(self):
        section = self.doc.sections[0]
        section.page_width = Cm(21.0)
        section.page_height = Cm(29.7)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.header_distance = Cm(1.25)
        section.footer_distance = Cm(1.25)

    def _setup_styles(self):
        styles = self.doc.styles
        normal = styles["Normal"]
        normal.font.name = "Times New Roman"
        normal.font.size = Pt(12)
        normal.font.color.rgb = DARK
        pf = normal.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_after = Pt(8)
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        for i, size, space_before in ((1, 16, 18), (2, 14, 14), (3, 12, 10)):
            style = styles[f"Heading {i}"]
            style.font.name = "Times New Roman"
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = NAVY
            style.paragraph_format.space_before = Pt(space_before)
            style.paragraph_format.space_after = Pt(8)
            style.paragraph_format.line_spacing = 1.15
            style.paragraph_format.keep_with_next = True

    def _setup_header_footer(self):
        section = self.doc.sections[0]
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = hp.add_run("Ubuzima Bwiza  |  Project Defense Report")
        set_run_font(run, size=9, italic=True, color=TEAL)

        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = fp.add_run("Page ")
        set_run_font(r1, size=10, color=GRAY)
        add_page_number(fp)
        r2 = fp.add_run("  ·  Confidential — for academic examination only")
        set_run_font(r2, size=9, italic=True, color=GRAY)

    def page_break(self):
        self.doc.add_page_break()

    def h(self, text: str, level: int = 1):
        self.doc.add_heading(text, level=level)

    def p(self, text: str, *, bold=False, italic=False, center=False, size=12, space_after=8, first_line=True):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        para.paragraph_format.space_after = Pt(space_after)
        if first_line and not center:
            para.paragraph_format.first_line_indent = Cm(1.25)
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic)
        return para

    def bullet(self, text: str, level: int = 0):
        para = self.doc.add_paragraph(style="List Bullet")
        para.clear()
        para.paragraph_format.left_indent = Cm(1.25 + level * 0.75)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        set_run_font(run, size=12)
        return para

    def numbered(self, text: str):
        para = self.doc.add_paragraph(style="List Number")
        para.clear()
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.space_after = Pt(4)
        run = para.add_run(text)
        set_run_font(run, size=12)
        return para

    def quote(self, text: str):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.paragraph_format.left_indent = Cm(1.5)
        para.paragraph_format.right_indent = Cm(1.0)
        para.paragraph_format.first_line_indent = Cm(0)
        para.paragraph_format.space_before = Pt(6)
        para.paragraph_format.space_after = Pt(10)
        run = para.add_run(text)
        set_run_font(run, size=11, italic=True, color=TEAL)
        return para

    def caption(self, text: str):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_before = Pt(4)
        para.paragraph_format.space_after = Pt(12)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        set_run_font(run, size=10, italic=True, color=GRAY)
        return para

    def table(self, headers: list[str], rows: list[list[str]], caption: str | None = None):
        tbl = self.doc.add_table(rows=1 + len(rows), cols=len(headers))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(h)
            set_run_font(run, size=10, bold=True, color=WHITE)
            shade_cell(cell, HEADER_BG)
            set_cell_borders(cell)
        for r_i, row in enumerate(rows):
            for c_i, val in enumerate(row):
                cell = tbl.rows[r_i + 1].cells[c_i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(val)
                set_run_font(run, size=10)
                if r_i % 2 == 1:
                    shade_cell(cell, ALT_BG)
                set_cell_borders(cell)
        prevent_table_split(tbl)
        if caption:
            self.caption(caption)
        else:
            self.doc.add_paragraph()

    def cover_line(self, text, size=14, bold=False, italic=False, space_after=6, color=DARK):
        para = self.doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(space_after)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(text)
        set_run_font(run, size=size, bold=bold, italic=italic, color=color)
        return para

    def save(self, path: Path):
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))


def build() -> Document:
    r = Report()

    # ------------------------------------------------------------------
    # COVER
    # ------------------------------------------------------------------
    for _ in range(2):
        r.cover_line("", size=12, space_after=6)
    r.cover_line("REPUBLIC OF RWANDA", size=13, bold=True, color=NAVY, space_after=2)
    r.cover_line("[NAME OF UNIVERSITY / INSTITUTION]", size=13, bold=True, color=NAVY, space_after=2)
    r.cover_line("[FACULTY / SCHOOL]", size=12, space_after=2)
    r.cover_line("[DEPARTMENT OF COMPUTER SCIENCE / INFORMATION TECHNOLOGY]", size=12, space_after=18)

    r.cover_line("A PROJECT REPORT SUBMITTED IN PARTIAL FULFILMENT", size=11, italic=True, space_after=2)
    r.cover_line("OF THE REQUIREMENTS FOR THE AWARD OF THE DEGREE OF", size=11, italic=True, space_after=8)
    r.cover_line("[Bachelor of Science in Computer Science / Information Technology]", size=12, bold=True, space_after=22)

    r.cover_line("PROJECT TITLE", size=11, bold=True, color=TEAL, space_after=8)
    r.cover_line("UBUZIMA BWIZA", size=26, bold=True, color=NAVY, space_after=6)
    r.cover_line(
        "A Multilingual Digital Health Platform for Specialist Discovery,",
        size=13,
        italic=True,
        space_after=2,
        color=NAVY,
    )
    r.cover_line(
        "Appointment Booking, Care Coordination and AI-Assisted",
        size=13,
        italic=True,
        space_after=2,
        color=NAVY,
    )
    r.cover_line("Symptom Triage in Rwanda", size=13, italic=True, space_after=22, color=NAVY)

    r.cover_line("Submitted by", size=11, space_after=4)
    r.cover_line("[STUDENT FULL NAME]", size=14, bold=True, space_after=2)
    r.cover_line("Registration No.: [REGISTRATION NUMBER]", size=12, space_after=16)

    r.cover_line("Supervisor", size=11, space_after=4)
    r.cover_line("[SUPERVISOR FULL NAME, TITLE]", size=13, bold=True, space_after=2)
    r.cover_line("[Academic rank / Department]", size=12, space_after=18)

    r.cover_line("September 2026", size=12, bold=True, space_after=4)
    r.cover_line("Kigali, Rwanda", size=12, italic=True, space_after=4)

    # ------------------------------------------------------------------
    # DECLARATION
    # ------------------------------------------------------------------
    r.page_break()
    r.h("DECLARATION")
    r.p(
        "I, [STUDENT FULL NAME], hereby declare that this project report entitled "
        "“Ubuzima Bwiza: A Multilingual Digital Health Platform for Specialist Discovery, "
        "Appointment Booking, Care Coordination and AI-Assisted Symptom Triage in Rwanda” "
        "is my original work. It has been prepared under the supervision of "
        "[SUPERVISOR FULL NAME] and has not been submitted, in whole or in part, for the award "
        "of any other degree, diploma or certificate in this or any other institution of higher learning."
    )
    r.p(
        "I further declare that all sources of information used in the preparation of this report "
        "have been duly acknowledged. Any limitation of the system as a working prototype — "
        "including the use of demonstration authentication and client-side persistence — is "
        "stated honestly in the body of this report."
    )
    r.p("Student signature: ________________________     Date: ______________", first_line=False)
    r.p("Name: [STUDENT FULL NAME]", first_line=False)
    r.p("Registration number: [REGISTRATION NUMBER]", first_line=False)

    r.h("SUPERVISOR’S APPROVAL", 2)
    r.p(
        "This project report has been submitted for examination with my approval as the "
        "university supervisor."
    )
    r.p("Supervisor signature: ________________________     Date: ______________", first_line=False)
    r.p("Name: [SUPERVISOR FULL NAME]", first_line=False)
    r.p("Title / Department: [TITLE]", first_line=False)

    # ------------------------------------------------------------------
    # ACKNOWLEDGEMENT
    # ------------------------------------------------------------------
    r.page_break()
    r.h("ACKNOWLEDGEMENT")
    r.p(
        "I thank the Almighty God for the strength and discipline required to complete this work. "
        "I express my sincere gratitude to my supervisor, [SUPERVISOR FULL NAME], for academic "
        "guidance, constructive criticism and encouragement throughout the design, implementation "
        "and documentation of Ubuzima Bwiza."
    )
    r.p(
        "I am grateful to the lecturers and administration of [UNIVERSITY / DEPARTMENT] for the "
        "knowledge and laboratory environment that made this project possible. I also thank my "
        "family and classmates for their support during development and testing."
    )
    r.p(
        "Any remaining errors of fact or interpretation are my own."
    )

    # ------------------------------------------------------------------
    # ABSTRACT
    # ------------------------------------------------------------------
    r.page_break()
    r.h("ABSTRACT")
    r.p(
        "Access to timely and trusted healthcare in Rwanda is still constrained by distance from "
        "specialist facilities, language barriers, fragmented appointment processes, and the "
        "absence of a single digital place where a patient can discover a clinician, book a visit, "
        "pay a consultation fee, keep a simple health record, and receive preliminary guidance "
        "before travelling. This project presents Ubuzima Bwiza (“Good Health” in Kinyarwanda), "
        "a multilingual web-based digital health platform designed to close that gap for patients, "
        "doctors, hospitals and platform administrators."
    )
    r.p(
        "The system enables a patient to search and compare specialists, book an in-person or "
        "video consultation, pay using locally relevant methods including MTN Mobile Money and "
        "Airtel Money (in Rwandan Francs), manage medications and records, apply for a chronic "
        "care programme, and converse with an AI Health Assistant that maps symptoms to a "
        "specialty and suggests bookable doctors. Doctors can approve, reject or complete "
        "appointments and review their calendar and patient list. Hospital operators oversee "
        "doctors, patients, reception appointments and operational reports. Platform administrators "
        "manage users, approvals, announcements and a trilingual content management facility."
    )
    r.p(
        "The platform was implemented as a responsive Single-Page Application using React 19, "
        "TypeScript and Vite, with full interface translation in English, Kinyarwanda and French. "
        "The distinctive technical contribution is a hybrid AI assistant: a built-in rule-based "
        "triage engine that works immediately without any external application programming interface "
        "key, together with an optional large-language-model pathway (Groq or OpenAI) for richer "
        "conversation. Emergency language is detected and the user is directed to urgent care. "
        "On every response the assistant states that the guidance is preliminary and is not a "
        "medical diagnosis. The same AI logic is deployed through a Vite development middleware, "
        "a Netlify Function and a Vercel serverless function, so the prototype is portable."
    )
    r.p(
        "The result is a complete, demonstrable end-to-end care journey tailored to the Rwandan "
        "context. As an academic prototype, authentication and most business data are persisted "
        "in the browser for examination and demonstration; production deployment would require a "
        "persistent backend, verified provider onboarding, live payment gateways and clinical "
        "governance. Within the scope of a final-year project, Ubuzima Bwiza is a coherent, "
        "locally grounded and technically non-trivial solution that is ready for oral defence."
    )
    r.p(
        "Keywords: digital health; telemedicine; AI symptom triage; appointment booking; "
        "mobile money; Kinyarwanda; Rwanda; multi-role web application.",
        italic=True,
        first_line=False,
    )

    # ------------------------------------------------------------------
    # TOC
    # ------------------------------------------------------------------
    r.page_break()
    r.h("TABLE OF CONTENTS")
    toc = [
        "Declaration",
        "Supervisor’s approval",
        "Acknowledgement",
        "Abstract",
        "List of tables",
        "List of abbreviations",
        "Chapter 1  Introduction",
        "     1.1  Background of the study",
        "     1.2  Problem statement",
        "     1.3  Research questions",
        "     1.4  General objective",
        "     1.5  Specific objectives",
        "     1.6  Scope of the study",
        "     1.7  Significance of the study",
        "     1.8  Limitations of the study",
        "     1.9  Organisation of the report",
        "Chapter 2  Literature review and related systems",
        "     2.1  Digital health and telemedicine",
        "     2.2  The Rwandan healthcare access context",
        "     2.3  AI in clinical triage and safety constraints",
        "     2.4  Related systems and the gap addressed",
        "Chapter 3  Methodology",
        "     3.1  Research and development approach",
        "     3.2  Requirements elicitation",
        "     3.3  Tools and technologies",
        "     3.4  Evaluation method",
        "Chapter 4  System analysis and design",
        "     4.1  Stakeholders and use cases",
        "     4.2  Functional requirements",
        "     4.3  Non-functional requirements",
        "     4.4  Architecture",
        "     4.5  Data model",
        "     4.6  Appointment and payment lifecycle",
        "     4.7  AI Health Assistant design",
        "Chapter 5  Implementation",
        "     5.1  Project structure and modules",
        "     5.2  Authentication and authorisation",
        "     5.3  Internationalisation and content management",
        "     5.4  Role dashboards and reporting",
        "     5.5  Deployment strategy",
        "Chapter 6  Testing, results and discussion",
        "     6.1  Test strategy",
        "     6.2  Results against objectives",
        "     6.3  Discussion",
        "Chapter 7  Conclusion and recommendations",
        "     7.1  Conclusion",
        "     7.2  Recommendations and future work",
        "References",
        "Appendix A  Demonstration accounts and how to run the system",
        "Appendix B  Route map of the application",
        "Appendix C  Suggested oral-defence questions and answers",
        "Appendix D  Fill-in cover details (student checklist)",
    ]
    for item in toc:
        para = r.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.line_spacing = 1.15
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(item)
        set_run_font(run, size=12)

    r.h("LIST OF TABLES", 2)
    for t in [
        "Table 1.  Stakeholder roles and primary capabilities",
        "Table 2.  Technology stack",
        "Table 3.  Functional requirements (summary)",
        "Table 4.  Non-functional requirements",
        "Table 5.  Local storage keys used by the prototype",
        "Table 6.  AI specialty detection rules (summary)",
        "Table 7.  Mapping of specific objectives to implemented features",
        "Table 8.  Demonstration login accounts",
    ]:
        para = r.doc.add_paragraph()
        para.paragraph_format.space_after = Pt(2)
        para.paragraph_format.first_line_indent = Cm(0)
        run = para.add_run(t)
        set_run_font(run, size=12)

    r.h("LIST OF ABBREVIATIONS", 2)
    r.table(
        ["Abbreviation", "Meaning"],
        [
            ["AI", "Artificial Intelligence"],
            ["API", "Application Programming Interface"],
            ["CHUK", "University Teaching Hospital of Kigali (Centre Hospitalier Universitaire de Kigali)"],
            ["CMS", "Content Management System"],
            ["CSV", "Comma-Separated Values"],
            ["EHR", "Electronic Health Record"],
            ["LLM", "Large Language Model"],
            ["MoMo", "Mobile Money (MTN)"],
            ["PDF", "Portable Document Format"],
            ["RWF", "Rwandan Franc"],
            ["SPA", "Single-Page Application"],
            ["UI", "User Interface"],
            ["UX", "User Experience"],
            ["WHO", "World Health Organization"],
        ],
    )

    # ------------------------------------------------------------------
    # CHAPTER 1
    # ------------------------------------------------------------------
    r.page_break()
    r.h("CHAPTER 1  INTRODUCTION")

    r.h("1.1  Background of the study", 2)
    r.p(
        "Healthcare delivery in Rwanda has improved substantially over the last two decades through "
        "community health workers, expanded insurance coverage and investment in referral hospitals. "
        "Nevertheless, a citizen who needs a specialist still faces a practical sequence of obstacles: "
        "knowing which specialty is appropriate, finding a trusted clinician, travelling to a city "
        "facility, obtaining an appointment by telephone or in person, paying a fee, and keeping "
        "track of prescriptions and follow-up dates. These steps are often disconnected. Language "
        "adds a further barrier. Many digital tools are published only in English or French, while "
        "a large share of the population is most comfortable in Kinyarwanda."
    )
    r.p(
        "Mobile penetration and mobile-money usage in Rwanda create an opportunity. A web platform "
        "that speaks the three official working languages of public life, that prices care in RWF, "
        "that accepts MTN Mobile Money and Airtel Money as first-class payment options, and that "
        "offers both in-person and video visits, can reduce unnecessary travel and waiting. "
        "A carefully constrained AI assistant can further help a patient describe symptoms and "
        "reach the right specialist faster — provided it never pretends to replace a clinician."
    )
    r.p(
        "Ubuzima Bwiza was therefore conceived as a unified digital front door to care: one product "
        "in which patients, doctors, hospital reception and a platform administrator each have a "
        "role-appropriate workspace. The Kinyarwanda name (“Good Health”) signals that the product "
        "is intended for every profession — farmers, constructors, artists and clinicians — not only "
        "for an urban elite."
    )
    r.quote(
        "“Find the right specialist, compare options and schedule your visit in just a few clicks "
        "wherever you are in Rwanda.”  — Ubuzima Bwiza product statement"
    )

    r.h("1.2  Problem statement", 2)
    r.p(
        "The core problem addressed by this project is the fragmentation of the outpatient care "
        "journey in Rwanda. Patients lack a single, language-inclusive digital channel to (i) "
        "discover verified specialists and hospitals, (ii) book in-person or telemedicine visits, "
        "(iii) pay consultation fees through locally used methods, (iv) retain a simple record of "
        "appointments, medications and documents, and (v) obtain preliminary, safety-aware guidance "
        "when they are unsure which specialty to consult. On the supply side, doctors and hospitals "
        "lack a lightweight workspace to receive booking requests, approve or reject them, and "
        "produce operational reports. Existing international telehealth products are poorly "
        "localised (language, currency, mobile money, facility names and emergency routing) and "
        "are rarely demonstrated as a complete multi-role system in an academic setting."
    )
    r.p(
        "Without a coherent digital layer, the cost of access remains hidden in travel time, missed "
        "work, delayed chronic-disease follow-up, and avoidable crowding of emergency units by "
        "patients who needed a scheduled specialist visit. This project treats that fragmentation "
        "as a software-engineering problem with a localised, demonstrable solution."
    )

    r.h("1.3  Research questions", 2)
    r.p("The study was guided by the following questions:", first_line=False)
    r.numbered(
        "How can a single web application serve four distinct healthcare stakeholders (patient, "
        "doctor, hospital, administrator) without confusing the user or mixing privileges?"
    )
    r.numbered(
        "How can specialist discovery, booking, local payment, records and messaging be joined "
        "into one continuous journey that a non-technical user can complete?"
    )
    r.numbered(
        "How can an AI symptom assistant be designed so that it is useful in Kinyarwanda and "
        "English, maps symptoms to specialties, suggests real doctors on the platform, and "
        "still refuses to issue a diagnosis?"
    )
    r.numbered(
        "How can the assistant remain usable when no commercial LLM key is available, yet upgrade "
        "to a live model when a key is configured?"
    )
    r.numbered(
        "What architecture allows the same AI logic to run in local development and on common "
        "student-friendly hosts (Netlify, Vercel)?"
    )

    r.h("1.4  General objective", 2)
    r.p(
        "To design, implement and demonstrate a multilingual digital health web platform — "
        "Ubuzima Bwiza — that localises specialist discovery, appointment booking, care "
        "coordination and AI-assisted symptom triage to the Rwandan context, and that is "
        "sufficiently complete for academic examination and oral defence."
    )

    r.h("1.5  Specific objectives", 2)
    r.p("The general objective was broken down as follows:", first_line=False)
    r.numbered(
        "To analyse the outpatient access problem and specify functional and non-functional "
        "requirements for patients, doctors, hospitals and administrators."
    )
    r.numbered(
        "To design a role-based architecture, data model and user interface that supports the "
        "full appointment lifecycle (book, pay, approve, reject, complete) together with "
        "medications, records, messaging, chronic-care application and operational reports."
    )
    r.numbered(
        "To implement the platform as a TypeScript React application with English, Kinyarwanda "
        "and French, light/dark theme, and Rwanda-specific content (facilities, RWF, mobile money, "
        "Africa/Kigali, local telephone format)."
    )
    r.numbered(
        "To implement a hybrid AI Health Assistant that performs rule-based triage without an "
        "API key, optionally calls Groq or OpenAI, detects emergencies, recommends a specialty "
        "and matching doctors, and always displays a medical disclaimer."
    )
    r.numbered(
        "To provide an administrator content-management facility so that marketing and legal "
        "strings can be edited in all three languages without redeploying code."
    )
    r.numbered(
        "To package the system so that it can be run locally for defence (npm run dev) and "
        "deployed to Netlify or Vercel with the AI endpoint intact."
    )
    r.numbered(
        "To test the main user journeys, map results to the objectives, and state honestly what "
        "must be added before production clinical use."
    )

    r.h("1.6  Scope of the study", 2)
    r.p("In scope.", first_line=False, bold=True)
    r.bullet("Public marketing site: home, doctor directory and profiles, booking, about, contact, privacy, terms, chronic-care information, public AI assistant.")
    r.bullet("Authentication screens: login, registration (patient, doctor, hospital), forgot password (demonstration).")
    r.bullet("Patient workspace: appointments, payments and receipts, messages, medications, medical records, chronic-care application, profile, AI chat.")
    r.bullet("Doctor workspace: dashboard, appointment actions, calendar, patients, availability, reports, messages, profile.")
    r.bullet("Hospital workspace: overview, doctors, patients, reception appointments, messages, CSV/PDF reports, settings.")
    r.bullet("Administrator workspace: users, doctor/hospital approvals, all appointments, payment overview, announcements, trilingual CMS, settings.")
    r.bullet("Hybrid AI triage with doctor-booking deep links; floating AI widget on all pages.")
    r.p("Out of scope (explicitly deferred to future work).", first_line=False, bold=True)
    r.bullet("A production database, server-side session security, and password hashing at rest.")
    r.bullet("Live MTN/Airtel payment-gateway settlement and reconciliation.")
    r.bullet("Real-time video consultation (WebRTC); video is modelled as a visit type.")
    r.bullet("Integration with national EHR, RSSB/Mutuelle claims, or licensed e-prescribing.")
    r.bullet("Medical-device connectivity and diagnostic imaging archives.")
    r.bullet("Full clinical validation of the triage rules by a licensed medical board.")
    r.p(
        "The system is therefore a high-fidelity academic prototype: the interfaces, workflows "
        "and AI behaviour can be demonstrated end to end, while persistence and payments are "
        "simulated in a form that a supervisor can inspect on a laptop."
    )

    r.h("1.7  Significance of the study", 2)
    r.p(
        "Academic significance. The project shows that a final-year student can combine modern "
        "front-end engineering, internationalisation, role-based access, serverless functions and "
        "constrained AI into one examinable artefact. It is not a single-page mock-up; it is a "
        "multi-workspace product with a documented architecture."
    )
    r.p(
        "Practical significance. If later industrialised, the design would reduce travel for "
        "routine specialist booking, support chronic-disease follow-up, and give hospital "
        "reception a shared appointment list. Mobile-money-first payment and Kinyarwanda as a "
        "first-class language are not cosmetic: they are the conditions for adoption outside "
        "central Kigali."
    )
    r.p(
        "Social significance. The product narrative explicitly includes farmers and other "
        "non-office workers. The AI assistant accepts a Kinyarwanda example prompt "
        "(“Ndashoje umutwe kandi mfite umuriro” — headache with fever) so that language is "
        "treated as a clinical-access issue, not only a translation task."
    )
    r.p(
        "Ethical significance. The assistant is instructed that it is not a doctor, must not "
        "diagnose with certainty, must escalate emergencies, and must attach a disclaimer. "
        "This is a deliberate safety property, not an afterthought, and is a point the candidate "
        "can defend."
    )

    r.h("1.8  Limitations of the study", 2)
    r.p(
        "Honesty about limitations is part of academic integrity and will strengthen, not weaken, "
        "the defence. The following constraints are known:"
    )
    r.bullet(
        "Demonstration authentication. Credentials are checked against seeded demo users. "
        "Newly registered accounts exist for the session but are not a durable user registry."
    )
    r.bullet(
        "Client-side persistence. The logged-in user and appointments are stored in browser "
        "localStorage. This is adequate for a viva demonstration on one machine; it is not a "
        "multi-user production database."
    )
    r.bullet(
        "Simulated payments and some messaging. Mobile-money and card flows issue a local receipt; "
        "they do not move real funds. Non-AI message threads use a short simulated reply delay."
    )
    r.bullet(
        "Rule-based triage is not a substitute for a licensed clinical protocol. Specialty mapping "
        "uses regular expressions and a fixed doctor catalogue. An LLM, when enabled, still cannot "
        "legally or ethically diagnose."
    )
    r.bullet(
        "No formal usability trial with a statistically sampled rural population was conducted "
        "within the project calendar; evaluation is based on walkthroughs of the implemented journeys."
    )

    r.h("1.9  Organisation of the report", 2)
    r.p(
        "Chapter 2 reviews digital health, the Rwandan access context, AI triage safety and related "
        "systems. Chapter 3 describes the methodology and tools. Chapter 4 presents analysis and "
        "design. Chapter 5 reports implementation. Chapter 6 presents testing, results and "
        "discussion. Chapter 7 concludes and recommends future work. Appendices contain run "
        "instructions, the route map, and suggested viva questions."
    )

    # ------------------------------------------------------------------
    # CHAPTER 2
    # ------------------------------------------------------------------
    r.page_break()
    r.h("CHAPTER 2  LITERATURE REVIEW AND RELATED SYSTEMS")

    r.h("2.1  Digital health and telemedicine", 2)
    r.p(
        "The World Health Organization describes digital health as the use of information and "
        "communication technologies to improve health, from electronic records to mobile health "
        "and telemedicine. Telemedicine — the delivery of clinical services at a distance — is "
        "particularly relevant where specialist density is low outside capital cities. Typical "
        "building blocks of a consumer telehealth product are: provider directory, scheduling, "
        "secure messaging, optional video, payment, and a personal health locker."
    )
    r.p(
        "International platforms (for example generalist telehealth marketplaces) prove that the "
        "booking-plus-payment pattern is viable. They also show a recurring failure mode when "
        "copied without localisation: English-only interfaces, card-only payments, specialty "
        "taxonomies that do not match local training pathways, and emergency advice that cites "
        "the wrong emergency number. A Rwandan-facing system must therefore treat language, "
        "currency, mobile money and local facilities as first-class requirements."
    )

    r.h("2.2  The Rwandan healthcare access context", 2)
    r.p(
        "Rwanda’s health system is organised from community health workers through health centres "
        "and district hospitals to national referral hospitals such as CHUK and specialised "
        "private facilities. Referral often implies travel to Kigali. For a parent in Musanze "
        "with a febrile child, or a farmer managing hypertension, the cost of that travel is not "
        "only the bus fare; it is lost labour and uncertainty about whether a specialist will be "
        "available on arrival."
    )
    r.p(
        "Kinyarwanda, English and French coexist in education, administration and clinical "
        "encounters. A digital product that omits Kinyarwanda excludes a large share of potential "
        "users. Equally, pricing in RWF and offering MTN MoMo and Airtel Money matches how "
        "households actually pay. Ubuzima Bwiza encodes these facts in the interface: hospital "
        "names familiar to Kigali users, Africa/Kigali as the operational timezone, and a "
        "+250 telephone pattern in contact information."
    )

    r.h("2.3  AI in clinical triage and safety constraints", 2)
    r.p(
        "Symptom checkers and large language models can reduce uncertainty about “which doctor "
        "should I see?”, but they introduce well-documented risks: hallucinated advice, missed "
        "emergencies, and the appearance of medical authority. Responsible design therefore "
        "imposes hard constraints: the system must identify itself as non-clinical; it must "
        "escalate red-flag symptoms (chest pain, severe breathing difficulty, stroke signs, "
        "unconsciousness, severe bleeding, and related emergencies) to immediate in-person care; "
        "and it must keep answers short and actionable rather than encyclopaedic."
    )
    r.p(
        "This project adopts a hybrid strategy that is appropriate for a student environment. "
        "A deterministic rule engine always works, even offline from commercial APIs. An optional "
        "LLM (Groq or OpenAI) can be switched on for more natural dialogue. If the live model "
        "fails, the system falls back to the rule engine and warns the user. The system prompt "
        "used for the LLM states:"
    )
    r.quote(
        "“You are Ubuzima Bwiza's AI Health Assistant for patients in Rwanda. You are NOT a "
        "doctor and must not diagnose with certainty. Always include a short disclaimer that "
        "this is preliminary guidance only.”"
    )
    r.p(
        "The assistant may reply in English or Kinyarwanda according to the user’s language. "
        "This is both a usability feature and a safety feature: people describe pain most "
        "accurately in the language they think in."
    )

    r.h("2.4  Related systems and the gap addressed", 2)
    r.p(
        "Related categories include hospital patient portals, international telehealth apps, "
        "generic clinic booking widgets, and standalone symptom-checker chatbots. Each category "
        "usually optimises for one actor. A hospital portal assumes the patient already belongs "
        "to that hospital. A chatbot may give advice but cannot complete a booking with a named "
        "local doctor. A booking widget may schedule a slot but offers no triage and no hospital "
        "operations view."
    )
    r.p(
        "The gap that Ubuzima Bwiza fills, within academic scope, is integration: one localised "
        "product in which triage output is not a dead-end paragraph but a bridge to a bookable "
        "specialist; in which payment methods match Rwanda; in which doctors and hospitals see "
        "the same appointment objects that patients create; and in which an administrator can "
        "edit public content in three languages. That integration is the intellectual centre of "
        "the project and the strongest argument for a passing defence."
    )

    # ------------------------------------------------------------------
    # CHAPTER 3
    # ------------------------------------------------------------------
    r.page_break()
    r.h("CHAPTER 3  METHODOLOGY")

    r.h("3.1  Research and development approach", 2)
    r.p(
        "The project followed an applied software-engineering methodology rather than a purely "
        "theoretical study. Requirements were derived from the documented access problem, from "
        "walkthroughs of a typical outpatient journey, and from the need to examine four roles "
        "in a single viva. Development used iterative prototyping: a public marketing shell, "
        "then authentication and role dashboards, then payments and reports, then the AI "
        "assistant and the administrator CMS."
    )
    r.p(
        "This approach is appropriate because the success criterion is a working demonstration, "
        "not a population survey. Where medical content appears (triage tips, disclaimers), the "
        "method was conservative: short, general, and always labelled as non-diagnostic."
    )

    r.h("3.2  Requirements elicitation", 2)
    r.p(
        "Requirements were elicited by decomposing the care journey into stages (discover, "
        "decide, book, pay, attend or connect, follow up) and by listing what each role must "
        "see at each stage. Constraints were added from the Rwandan context (languages, RWF, "
        "mobile money, known Kigali facilities) and from examination practicality (the system "
        "must boot with npm run dev and remain usable if no LLM key is present)."
    )

    r.h("3.3  Tools and technologies", 2)
    r.p(
        "The stack was chosen for type safety, rapid user-interface iteration, and free hosting "
        "options familiar to students."
    )
    r.table(
        ["Layer", "Choice", "Rationale"],
        [
            ["User interface", "React 19 + TypeScript", "Component model for four role shells; compile-time checks"],
            ["Build tool", "Vite 8", "Fast local server; plugin hook for the AI API in development"],
            ["Routing", "React Router DOM v7", "Nested public and protected routes"],
            ["Languages", "i18next + react-i18next", "English, Kinyarwanda, French with persisted locale"],
            ["Styling", "Custom CSS, CSS variables", "Light / dark / system theme without a heavy UI kit"],
            ["State", "React Context", "Auth, CMS overrides, theme — sufficient at prototype scale"],
            ["AI logic", "Shared server/aiChat.ts", "One implementation for Vite, Netlify and Vercel"],
            ["Optional LLM", "Groq or OpenAI HTTP APIs", "Upgrade path without rewriting triage"],
            ["Lint", "oxlint", "Fast static checks"],
            ["Hosting targets", "Netlify, Vercel, GitHub Pages", "SPA fallback plus serverless AI where supported"],
        ],
        "Table 2. Technology stack.",
    )
    r.p(
        "No traditional Express server and no SQL database were introduced, by design. The "
        "academic goal was a complete interactive product that a supervisor can run on a "
        "workstation. Introducing an under-tested backend would have expanded risk without "
        "improving the demonstration of localisation, roles and AI. The report is explicit "
        "that a production sequel would add those layers."
    )

    r.h("3.4  Evaluation method", 2)
    r.p(
        "Evaluation used structured walkthroughs against the specific objectives: each role was "
        "logged in with the published demonstration accounts; each critical path (search–book–pay, "
        "doctor approval, hospital report download, admin CMS save, AI emergency phrase, AI "
        "specialty match, language switch) was executed. Results are reported in Chapter 6 as a "
        "traceability table from objectives to features. This is a verification-oriented method "
        "suitable for a prototype; it does not claim clinical efficacy."
    )

    # ------------------------------------------------------------------
    # CHAPTER 4
    # ------------------------------------------------------------------
    r.page_break()
    r.h("CHAPTER 4  SYSTEM ANALYSIS AND DESIGN")

    r.h("4.1  Stakeholders and use cases", 2)
    r.table(
        ["Role", "Primary goal", "Key use cases"],
        [
            [
                "Patient",
                "Obtain the right care with minimum travel and confusion",
                "Search doctors; book in-person or video; pay (MoMo, Airtel, card, cash); view medications and records; chat with AI; apply for chronic care; download reports",
            ],
            [
                "Doctor",
                "Control the appointment queue and follow patients",
                "See dashboard statistics; approve / reject / complete visits; calendar; patient list; availability; reports; messages",
            ],
            [
                "Hospital",
                "Coordinate reception and facility operations",
                "Overview; manage doctor and patient lists; reception appointments; messages; CSV and printable reports; facility settings (e.g. CHUK, Africa/Kigali)",
            ],
            [
                "Administrator",
                "Govern the platform and public content",
                "Users; doctor and hospital approvals; all appointments; payment overview; announcements; trilingual CMS (dozens of editable fields); theme and language settings",
            ],
        ],
        "Table 1. Stakeholder roles and primary capabilities.",
    )
    r.p(
        "Registration is offered for patient, doctor and hospital. The administrator account is "
        "seeded for examination so that governance features can be shown without an unsafe "
        "self-service promotion to admin."
    )

    r.h("4.2  Functional requirements", 2)
    r.table(
        ["ID", "Requirement", "Priority"],
        [
            ["FR-01", "Users shall log in by role and be redirected to the correct dashboard", "Must"],
            ["FR-02", "Patients shall filter the doctor directory by specialty, care type and name", "Must"],
            ["FR-03", "Patients shall book a dated time slot as in-person or video, with optional notes", "Must"],
            ["FR-04", "Patients shall pay an unpaid appointment and receive a receipt identifier", "Must"],
            ["FR-05", "Doctors shall change appointment status to approved, rejected or completed", "Must"],
            ["FR-06", "Hospital staff shall view doctors, patients and the full appointment list", "Must"],
            ["FR-07", "The AI assistant shall accept a symptom description and return guidance plus specialty", "Must"],
            ["FR-08", "The AI assistant shall detect emergency language and urge immediate in-person care", "Must"],
            ["FR-09", "The AI assistant shall suggest matching doctors with a path to booking", "Must"],
            ["FR-10", "The interface shall be switchable among English, Kinyarwanda and French", "Must"],
            ["FR-11", "Administrators shall edit selected public strings per language and persist overrides", "Should"],
            ["FR-12", "Users shall export appointment reports as CSV and as printable HTML/PDF", "Should"],
            ["FR-13", "Patients shall view medications and request AI medication tips", "Should"],
            ["FR-14", "Patients shall submit a chronic-care interest form", "Should"],
            ["FR-15", "A floating AI widget shall be available on all pages", "Should"],
        ],
        "Table 3. Functional requirements (summary).",
    )

    r.h("4.3  Non-functional requirements", 2)
    r.table(
        ["ID", "Requirement", "How addressed in the prototype"],
        [
            ["NFR-01", "Usability", "Role-specific navigation; demo-account buttons for examiners; clear CTAs"],
            ["NFR-02", "Localisation", "Three locale files; CMS overrides; AI language instruction"],
            ["NFR-03", "Availability of AI", "Rule engine works with zero configuration; LLM is optional"],
            ["NFR-04", "Portability", "Shared AI module; Netlify, Vercel and GitHub Pages configs"],
            ["NFR-05", "Safety messaging", "Disclaimer on AI replies; emergency escalation text"],
            ["NFR-06", "Theme accessibility", "Light, dark and system appearance"],
            ["NFR-07", "Examiner convenience", "Single command to run; documented demo passwords"],
            ["NFR-08", "Performance (interactive)", "SPA with Vite; last 12 chat messages sent to the API"],
        ],
        "Table 4. Non-functional requirements.",
    )

    r.h("4.4  Architecture", 2)
    r.p(
        "The runtime architecture is a browser Single-Page Application talking to one AI HTTP "
        "endpoint. Three Context providers wrap the tree: authentication and appointments, "
        "content overrides for the CMS, and theme. React Router distinguishes public marketing "
        "pages, authentication pages, and four protected dashboards. A shared RequireAuth guard "
        "sends anonymous users to login and sends a logged-in user of the wrong role back to "
        "their own home."
    )
    r.p(
        "The AI request path is POST /api/ai/chat with a short message history. In development, "
        "a Vite plugin serves that path. On Netlify, a redirect maps it to a Function. On Vercel, "
        "the api/ai/chat module serves it. All three call the same handleAiChat implementation. "
        "Provider preference is Groq, then OpenAI, then local triage. Empty environment keys are "
        "ignored so that a blank .env file cannot accidentally disable a real key."
    )
    r.p("Logical architecture (for the oral presentation):", first_line=False)
    r.bullet("Presentation layer: React pages and dashboard shell, i18n, theme.")
    r.bullet("Application layer: AuthContext (session, appointments, payments), ContentContext (CMS), AI client.")
    r.bullet("Intelligence layer: localAssistantReply rules + optional LLM with system prompt.")
    r.bullet("Integration adapters: Vite middleware, Netlify Function, Vercel API route.")
    r.bullet("Demonstration data layer: seeded doctors, users, appointments, medications, records; localStorage.")

    r.h("4.5  Data model", 2)
    r.p(
        "The principal types are User, Doctor, Appointment and supporting enumerations. A Role "
        "is one of patient, doctor, hospital or admin. An Appointment carries doctor identity, "
        "patient name, date and time, status (pending, approved, completed, rejected, cancelled), "
        "visit type (in-person or video), optional notes, amount in RWF, payment status, optional "
        "payment method (momo, airtel, card, cash), paid timestamp and receipt identifier."
    )
    r.table(
        ["Storage key", "Contents"],
        [
            ["ub_user", "Logged-in user profile (password is not stored)"],
            ["ub_appointments", "Shared appointment list visible across roles on the same browser"],
            ["ub_locale", "Selected language code (en, rw, fr)"],
            ["ub_theme", "light, dark or system"],
            ["ub_content_overrides", "Administrator CMS edits, keyed by locale"],
        ],
        "Table 5. Local storage keys used by the prototype.",
    )
    r.p(
        "Seed data includes four demonstration users (for example a patient named Aline Mukamana "
        "and Dr Jean Mugabo), a catalogue of doctors attached to Rwandan facilities such as CHUK "
        "and King Faisal Hospital, initial appointments in mixed statuses, sample medications "
        "(including common chronic-care drugs such as amlodipine and metformin), and sample "
        "records. This allows an examiner to see a populated system immediately after login."
    )

    r.h("4.6  Appointment and payment lifecycle", 2)
    r.numbered("The patient searches or opens a doctor profile and chooses Book.")
    r.numbered("Date, time slot (working hours in the interface), visit type and notes are captured. Optional AI visit-preparation notes can be generated.")
    r.numbered("The appointment is created as pending and unpaid.")
    r.numbered("The patient is directed to payment. Methods are MTN MoMo, Airtel Money, card, or cash at the facility.")
    r.numbered("On simulated success, paymentStatus becomes paid, a receiptId is issued, and a still-pending visit is auto-approved so that the clinical queue does not stall on a paid request.")
    r.numbered("The doctor may still reject or later mark the visit completed. Hospital reception sees the same objects. Reports aggregate paid totals, unpaid counts and the share of video visits.")

    r.h("4.7  AI Health Assistant design", 2)
    r.p(
        "Input is the recent conversation (capped at twelve messages). Output is a structured "
        "result: reply text, mode (local or llm), inferred specialty, zero or more doctor "
        "suggestions, and an optional warning if the live model was unavailable."
    )
    r.p(
        "Local mode inspects the last user utterance. If it is empty, the assistant asks for "
        "symptoms, age group and duration. If emergency patterns match, the reply insists on "
        "emergency services or the nearest hospital and does not encourage waiting for an "
        "online consult. Otherwise a specialty is inferred from ordered regular-expression rules, "
        "contextual first-aid-style tips are attached, up to two doctors of that specialty are "
        "selected from the server-side catalogue, and a disclaimer is appended."
    )
    r.table(
        ["Example user language (patterns)", "Suggested specialty"],
        [
            ["Chest pain, palpitations, blood pressure, hypertension", "Cardiologist"],
            ["Child, baby, infant, pediatric fever", "Pediatrician"],
            ["Pregnancy, menstrual, antenatal, women’s health", "Gynecologist"],
            ["Skin, rash, acne, eczema", "Dermatologist"],
            ["Tooth, gum, dental, cavity", "Dental"],
            ["Migraine, seizure, stroke, numbness", "Neurologist"],
            ["Eye, vision, blurry", "Ophthalmologist"],
            ["Bone, joint, fracture, back pain", "Orthopedic surgeon"],
            ["Diabetes, thyroid, hormone", "Endocrinologist"],
            ["Ear, nose, throat, sinus", "ENT surgeon"],
            ["Cough, cold, flu, fever, headache, stomach (default medical)", "General practitioner"],
        ],
        "Table 6. AI specialty detection rules (summary).",
    )
    r.p(
        "LLM mode uses a Rwanda-specific system prompt, the same doctor-matching idea, and the "
        "same disclaimer policy. Failure of the live API (other than a hard authentication error) "
        "returns local triage with the warning that live AI is temporarily unavailable. AI is "
        "surfaced in five places: the public /ai-assistant page, a floating widget, the messages "
        "inbox AI thread, booking-time visit preparation, and medication tips. This repetition is "
        "intentional: triage is useless if it is hidden in one menu."
    )

    # ------------------------------------------------------------------
    # CHAPTER 5
    # ------------------------------------------------------------------
    r.page_break()
    r.h("CHAPTER 5  IMPLEMENTATION")

    r.h("5.1  Project structure and modules", 2)
    r.p(
        "The source tree separates presentation (src/pages, src/components), shared demonstration "
        "data (src/data.ts), contexts, internationalisation JSON, AI client, report generation, "
        "and the portable AI server module. Netlify and Vercel wrappers are thin: they parse the "
        "HTTP request and call handleAiChat. This is a clean engineering choice; it shows the "
        "candidate can avoid duplicating business logic across hosts."
    )
    r.p("Principal implementation files the examiner may open:", first_line=False)
    r.bullet("src/App.tsx — route table for marketing, auth and the four workspaces.")
    r.bullet("src/context/AuthContext.tsx — session, appointment CRUD and payment updates.")
    r.bullet("src/components/dashboard/Shell.tsx — RequireAuth and per-role navigation.")
    r.bullet("src/components/AiChat.tsx — chat user interface, inbox integration, floating widget.")
    r.bullet("server/aiChat.ts — triage rules, LLM call, doctor matching.")
    r.bullet("src/i18n/locales/en.json, rw.json, fr.json — complete UI translation.")
    r.bullet("src/pages/admin/AdminContentPage.tsx — trilingual CMS.")
    r.bullet("src/lib/reports.ts — CSV and printable reports.")
    r.bullet("vite.config.ts, netlify.toml, vercel.json — run and deploy wiring.")

    r.h("5.2  Authentication and authorisation", 2)
    r.p(
        "Login compares email (case-insensitive) and password to the seeded directory. On success, "
        "the user object is kept in React state and in localStorage under ub_user. dashboardPath() "
        "maps role to the first protected page. Wrong-role access is rejected by the shell, which "
        "prevents a patient URL from revealing a doctor tool. Logout clears the session but leaves "
        "appointments in place so that a second role on the same browser can still see the shared "
        "queue — useful during a viva when the candidate switches accounts quickly."
    )
    r.p(
        "Registration allocates an identifier, appends the user to the in-memory directory and "
        "signs the person in. Google login is presented on the form but refused with a "
        "demonstration-only message, which is preferable to a silent fake social login."
    )

    r.h("5.3  Internationalisation and content management", 2)
    r.p(
        "i18next is initialised with English as fallback. The locale is taken from a stored "
        "preference or from the browser language when it starts with rw or fr. A language "
        "switcher is available in the chrome of the site. Specialty labels are translated through "
        "resource keys rather than hardcoded English in every card."
    )
    r.p(
        "The administrator CMS exposes a large set of editable fields (home, about, contact, "
        "legal, announcements, selected doctor bios) in each locale. Saves write to "
        "ub_content_overrides. This demonstrates a governance requirement: a platform operator "
        "must be able to correct public wording — including Kinyarwanda — without a developer "
        "deploy. Legal pages themselves are labelled as demonstration copy, which is the correct "
        "ethical stance before counsel reviews them."
    )

    r.h("5.4  Role dashboards and reporting", 2)
    r.p(
        "Each workspace reuses a dashboard shell (navigation, statistics, badges) so that the "
        "product feels like one system, not four unrelated sites. Patient home emphasises unpaid "
        "items and upcoming visits. Doctor home emphasises queue actions. Hospital home emphasises "
        "throughput. Admin home emphasises platform counts."
    )
    r.p(
        "Reports compute simple operational indicators: paid revenue, unpaid count, and video-visit "
        "ratio. Export is CSV for spreadsheets and a printable HTML view that the browser can save "
        "as PDF. Receipts follow the same printable path. For a prototype this is the right depth: "
        "the candidate can show a hospital administrator downloading evidence, without claiming a "
        "full business-intelligence warehouse."
    )

    r.h("5.5  Deployment strategy", 2)
    r.p(
        "Local defence: from the healthline directory, npm install then npm run dev, then open "
        "http://localhost:5173/. Optional live LLM: copy .env.example to .env, set GROQ_API_KEY "
        "or OPENAI_API_KEY, restart the server."
    )
    r.p(
        "Netlify: production build of the SPA, functions directory for AI, redirect from "
        "/api/ai/chat, and a catch-all redirect to index.html for client-side routes. Vercel: "
        "equivalent SPA rewrite with /api passthrough. GitHub Pages: a special build that sets "
        "the Vite base path and copies index.html to 404.html so that deep links degrade gracefully. "
        "Together these files show that the candidate thought about how a supervisor or an "
        "external examiner would actually open the work."
    )

    # ------------------------------------------------------------------
    # CHAPTER 6
    # ------------------------------------------------------------------
    r.page_break()
    r.h("CHAPTER 6  TESTING, RESULTS AND DISCUSSION")

    r.h("6.1  Test strategy", 2)
    r.p(
        "Testing combined (i) static TypeScript compilation and linting, (ii) manual end-to-end "
        "walkthroughs for each role, (iii) language-switch checks on representative screens, "
        "(iv) AI cases: empty input, ordinary symptom, Kinyarwanda example, and emergency phrase, "
        "and (v) negative tests: wrong password, wrong-role URL, non-POST to the AI function. "
        "Because the product is a user-interface prototype, manual scenario testing against the "
        "requirements table is more informative than a large unit-test count."
    )

    r.h("6.2  Results against objectives", 2)
    r.table(
        ["Specific objective", "Result in the implemented system"],
        [
            [
                "Analyse and specify multi-role requirements",
                "Achieved. Chapter 4 records FR/NFR; four shells match the specification.",
            ],
            [
                "Design lifecycle, records, messaging, chronic care, reports",
                "Achieved. Book–pay–approve/reject/complete is implemented; reports export CSV/PDF.",
            ],
            [
                "Implement trilingual React/TypeScript app with Rwanda localisation",
                "Achieved. en/rw/fr, RWF, MoMo/Airtel, local facilities, Kigali timezone in hospital settings.",
            ],
            [
                "Hybrid AI with emergency handling and doctor suggestions",
                "Achieved. Local mode always on; LLM optional; disclaimer present; booking links offered.",
            ],
            [
                "Administrator trilingual CMS",
                "Achieved. Overrides persist locally and flow through the public pages.",
            ],
            [
                "Runnable locally and packable for Netlify/Vercel",
                "Achieved. npm run dev; netlify.toml and vercel.json present with AI wiring.",
            ],
            [
                "Test journeys and state production gaps honestly",
                "Achieved. This chapter and section 1.8 document prototype limits.",
            ],
        ],
        "Table 7. Mapping of specific objectives to implemented features.",
    )
    r.p(
        "Observed behaviour during walkthroughs: demo buttons on the login page reduce examiner "
        "friction; paying an appointment updates the patient list and hospital reception view on "
        "the same browser; switching language updates navigation and home copy; the AI floating "
        "widget remains reachable from dashboards; an emergency-style phrase produces an urgent-care "
        "message rather than a casual booking suggestion. These observations constitute the "
        "empirical result of the project: the journeys work as designed."
    )

    r.h("6.3  Discussion", 2)
    r.p(
        "The strongest result is integration under localisation constraints. Many student projects "
        "either build a chatbot or build a CRUD dashboard. Ubuzima Bwiza connects the two and "
        "then adds hospital and admin governance. The hybrid AI decision is also defensible: a "
        "viva laboratory may have no paid API key, yet the candidate can still show triage. When "
        "a key is present, the same screens become conversational without a second product."
    )
    r.p(
        "The principal scientific and ethical caution is not to over-claim. The rule engine’s "
        "specialty mapping is a heuristic. It can misclassify ambiguous complaints. That is why "
        "the default specialty is a general practitioner and why the disclaimer is mandatory. "
        "A supervisor who asks “Is this medically validated?” should receive a direct answer: "
        "No. It is an engineering prototype of a triage-to-booking bridge, designed so that a "
        "future clinical team could replace the rule table with an approved protocol."
    )
    r.p(
        "A second caution concerns data. localStorage is visible to any script on the origin and "
        "is unsuitable for real patient secrets. The candidate should volunteer this point before "
        "being asked. Doing so demonstrates professional maturity and protects the project from "
        "the accusation that health data was treated casually."
    )
    r.p(
        "Overall, relative to the stated scope, the implementation meets the objectives. Relative "
        "to a national digital-health programme, it is a starting point. Both sentences can be "
        "true at once, and that is the correct academic position."
    )

    # ------------------------------------------------------------------
    # CHAPTER 7
    # ------------------------------------------------------------------
    r.page_break()
    r.h("CHAPTER 7  CONCLUSION AND RECOMMENDATIONS")

    r.h("7.1  Conclusion", 2)
    r.p(
        "This project set out to reduce the fragmentation of the outpatient journey in Rwanda by "
        "building Ubuzima Bwiza, a multilingual digital health platform. The implemented system "
        "allows a patient to find a specialist, book an in-person or video visit, pay using "
        "familiar local methods, keep a simple health locker, and obtain preliminary AI guidance "
        "that ends in a booking suggestion rather than a dead end. Doctors, hospitals and a "
        "platform administrator have corresponding workspaces. English, Kinyarwanda and French "
        "are supported throughout the interface, and the AI assistant is instructed to respect "
        "the user’s language while refusing to pose as a doctor."
    )
    r.p(
        "Technically, the work demonstrates a coherent Single-Page Application architecture, "
        "role-based routing, a portable hybrid AI module, and deployment descriptors for common "
        "hosts. Academically, the work is complete enough to run in an oral defence, to explain "
        "with diagrams, and to critique with a clear list of production next steps. The general "
        "and specific objectives stated in Chapter 1 have been achieved within the declared scope."
    )
    r.p(
        "Ubuzima Bwiza therefore stands as a locally grounded, technically non-trivial, and "
        "ethically cautious contribution to student digital-health engineering, and is submitted "
        "for examination with the recommendation that it proceeds to oral defence."
    )

    r.h("7.2  Recommendations and future work", 2)
    r.p("The following steps would convert the prototype into a service that could be piloted with a partner facility:", first_line=False)
    r.numbered(
        "Introduce a proper backend (for example a REST or GraphQL API) with a relational or "
        "document database, hashed passwords, server-side sessions or tokens, and audit logs."
    )
    r.numbered(
        "Integrate official MTN MoMo and Airtel Money collection APIs, with reconciliation and refunds."
    )
    r.numbered(
        "Replace demonstration provider records with a verification workflow (licence numbers, hospital affiliation, administrator approval that persists)."
    )
    r.numbered(
        "Add real-time video using WebRTC, with waiting-room and consent screens.")
    r.numbered(
        "Subject the triage rule table (and any LLM prompt) to review by licensed clinicians and the relevant ethics body; add a documented clinical-safety case.")
    r.numbered(
        "Integrate, where policy allows, with facility EHR or national digital-health building blocks rather than a private locker only.")
    r.numbered(
        "Conduct a supervised usability study with Kinyarwanda-dominant users outside Kigali, measuring task success for search, booking and AI explanation.")
    r.numbered(
        "Harden privacy: encryption at rest, role-based record access on the server, consent management, and a data-retention policy.")
    r.p(
        "Until those steps are complete, the system should be presented only as an academic "
        "demonstration, never as a clinical device or as a substitute for emergency services."
    )

    # ------------------------------------------------------------------
    # REFERENCES
    # ------------------------------------------------------------------
    r.page_break()
    r.h("REFERENCES")
    r.p(
        "The following sources informed the problem framing, safety stance and technical choices. "
        "The candidate should add any additional textbooks, lecture notes and URLs required by "
        "the department’s citation style (APA, IEEE or Harvard) before printing the bound copy.",
        first_line=False,
    )
    refs = [
        "World Health Organization. (2021). Global strategy on digital health 2020–2025. Geneva: WHO.",
        "World Health Organization. (2019). WHO guideline: Recommendations on digital interventions for health system strengthening. Geneva: WHO.",
        "Ministry of Health, Republic of Rwanda. (n.d.). Health sector policies and digital health programme documents. Kigali: Ministry of Health. [Update with the specific policy edition used in your literature review.]",
        "National Institute of Statistics of Rwanda. (n.d.). Demographic and health / ICT related publications as cited in the final printed copy.",
        "OpenAI. (2024–2026). API documentation for chat completions. https://platform.openai.com/docs",
        "Groq. (2024–2026). GroqCloud API documentation. https://console.groq.com/docs",
        "Meta Open Source. (2024). React documentation. https://react.dev",
        "Vite contributors. (2025–2026). Vite documentation. https://vite.dev",
        "React Router. (2025–2026). React Router documentation. https://reactrouter.com",
        "i18next. (n.d.). i18next documentation. https://www.i18next.com",
        "Netlify. (n.d.). Netlify Functions documentation. https://docs.netlify.com/functions/overview/",
        "Vercel. (n.d.). Vercel serverless functions documentation. https://vercel.com/docs/functions",
        "MDN Web Docs. (n.d.). Window.localStorage. https://developer.mozilla.org/en-US/docs/Web/API/Window/localStorage",
        "W3C. (n.d.). Web Content Accessibility Guidelines (WCAG) 2.2. https://www.w3.org/TR/WCAG22/",
        "Bates, D. W., & Gawande, A. A. (2003). Improving safety with information technology. New England Journal of Medicine, 348, 2526–2534. (Foundational argument for digital tools in care coordination; cite if your department expects journal sources.)",
        "Tsvetkova, M., et al. — replace this placeholder with any telemedicine or HCI papers actually read during the project, following departmental style.",
    ]
    for i, ref in enumerate(refs, 1):
        para = r.doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.25)
        para.paragraph_format.first_line_indent = Cm(-1.25)
        para.paragraph_format.space_after = Pt(6)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(f"[{i}]  {ref}")
        set_run_font(run, size=12)

    r.p(
        "Note to the candidate: examiners notice invented citations. Before binding, delete any "
        "placeholder line you did not actually consult, and add the exact MoH or university library "
        "items you used. The WHO digital-health strategy and the official framework documentation "
        "of React, Vite and your hosting provider are sufficient as starting technical references.",
        italic=True,
        first_line=False,
    )

    # ------------------------------------------------------------------
    # APPENDIX A
    # ------------------------------------------------------------------
    r.page_break()
    r.h("APPENDIX A  DEMONSTRATION ACCOUNTS AND HOW TO RUN THE SYSTEM")
    r.p(
        "These instructions allow the supervisor to reproduce the defence demonstration on a "
        "workstation with Node.js installed.",
        first_line=False,
    )
    r.h("A.1  Run locally", 2)
    r.p("In a terminal:", first_line=False)
    r.p("cd healthline", first_line=False, italic=True)
    r.p("npm install", first_line=False, italic=True)
    r.p("npm run dev", first_line=False, italic=True)
    r.p("Then open http://localhost:5173/ (or http://127.0.0.1:5173/).", first_line=False)
    r.p(
        "Optional live AI: copy .env.example to .env, set GROQ_API_KEY (recommended for a free/fast "
        "model such as llama-3.3-70b-versatile) or OPENAI_API_KEY, then restart npm run dev. "
        "Without a key, smart triage mode still works.",
        first_line=False,
    )

    r.h("A.2  Demonstration accounts", 2)
    r.table(
        ["Role", "Email", "Password"],
        [
            ["Patient", "patient@ubuzimabwiza.com", "patient123"],
            ["Doctor", "doctor@ubuzimabwiza.com", "doctor123"],
            ["Hospital", "hospital@ubuzimabwiza.com", "hospital123"],
            ["Admin", "admin@ubuzimabwiza.com", "admin123"],
        ],
        "Table 8. Demonstration login accounts.",
    )
    r.p(
        "The login page also provides one-click demo buttons. Recommended viva path: patient "
        "search and book; pay with MoMo; open AI assistant with a symptom and with an emergency "
        "phrase; switch to Kinyarwanda; log in as doctor and approve; log in as hospital and "
        "download a report; log in as admin and edit one CMS field.",
        first_line=False,
    )

    # ------------------------------------------------------------------
    # APPENDIX B
    # ------------------------------------------------------------------
    r.h("APPENDIX B  ROUTE MAP OF THE APPLICATION")
    r.p("Public marketing: /  /doctors  /doctors/:id  /book/:id  /about  /contact  /privacy-policy  /terms-of-service  /patient/chronic-care  /ai-assistant", first_line=False)
    r.p("Authentication: /login  /register  /forgot-password  (/signup redirects to /register)", first_line=False)
    r.p("Patient: /my-appointments  /payments  /pay/:id  /messages  /medications  /my-prescriptions  /medical-record  /records  /patient/chronic-care/apply  /my-profile", first_line=False)
    r.p("Doctor: /doctor-dashboard  /doctor-appointments  /doctor-calendar  /patients  /doctor/availability  /doctor/chronic-care  /doctor-profile  /doctor-reports  /doctor-messages", first_line=False)
    r.p("Hospital: /hospital-dashboard  /hospital-dashboard/doctors  /hospital-dashboard/patients  /hospital-dashboard/reception/appointments  /hospital-dashboard/messages  /hospital-dashboard/reports  /hospital-dashboard/settings", first_line=False)
    r.p("Admin: /admin-dashboard  /manage-users  /admin-content  /doctor-approvals  /hospital-approvals  /all-appointments  /payment-approvals  /announcements  /settings", first_line=False)

    # ------------------------------------------------------------------
    # APPENDIX C
    # ------------------------------------------------------------------
    r.page_break()
    r.h("APPENDIX C  SUGGESTED ORAL-DEFENCE QUESTIONS AND ANSWERS")
    r.p(
        "The candidate should rehearse these answers in their own words. They are written here "
        "to show the supervisor that the student can defend design decisions, not only click through screens.",
        first_line=False,
    )

    qa = [
        (
            "What problem does this project solve that a static hospital website does not?",
            "A static site can publish telephone numbers. Ubuzima Bwiza completes a journey: "
            "discover a specialist, book a slot, pay in RWF with mobile money, notify the clinical "
            "and hospital roles, and offer constrained AI triage that ends in a booking link. "
            "It also adds doctor, hospital and admin workspaces on the same data objects.",
        ),
        (
            "Why is this not a medical device, and how do you protect users from unsafe AI advice?",
            "The assistant is labelled as preliminary guidance, is forbidden to diagnose, escalates "
            "emergency phrases to in-person care, and falls back to a deterministic rule engine. "
            "It is an access and navigation aid, not a diagnostic device. Production use would still "
            "need clinical review and legal classification.",
        ),
        (
            "Why store data in localStorage if you know it is insecure for health records?",
            "Because the academic goal was a demonstrable multi-role product on one laptop. "
            "localStorage lets the examiner see persistence without standing up a database. "
            "I state this as a limitation and list a real backend, hashing and audit logs as future work. "
            "I would never store real patient records this way.",
        ),
        (
            "What is original about the AI component?",
            "The originality is the hybrid, portable, booking-aware design: rules plus optional LLM, "
            "shared code across Vite/Netlify/Vercel, specialty matching against the same doctor "
            "catalogue the UI uses, and Kinyarwanda/English behaviour. The algorithms themselves "
            "are transparent regular expressions plus a hosted model — I do not claim a new neural architecture.",
        ),
        (
            "How did you choose the technology stack?",
            "React and TypeScript for a complex UI with four roles; Vite for speed and a plugin "
            "that hosts the AI during development; i18next because localisation is a core requirement, "
            "not a plugin; serverless adapters so the same AI can be shown on free hosts. A heavy "
            "back end was postponed so that localisation, roles and AI could be finished to defence quality.",
        ),
        (
            "How does the system support users who prefer Kinyarwanda?",
            "The entire interface has a rw locale. The language switcher persists the choice. "
            "The CMS can override public strings in Kinyarwanda. The AI system prompt allows "
            "Kinyarwanda replies, and the assistant includes a Kinyarwanda quick prompt.",
        ),
        (
            "What happens if Groq or OpenAI is down during the viva?",
            "The product still works. Local triage mode does not need a key. If a key is configured "
            "and the live call fails, the user is warned and the rule engine answers.",
        ),
        (
            "How would you scale this to a real hospital?",
            "Move identity, appointments and payments to a server; integrate MoMo/Airtel officially; "
            "verify clinicians; add WebRTC; put triage under clinical governance; connect to the "
            "facility record system; run a usability study; apply privacy law and ethics approval.",
        ),
        (
            "What did you personally implement, and what would you improve first if you had four more weeks?",
            "Be ready to point at AuthContext, aiChat.ts, the dashboard shell, i18n files and the CMS. "
            "The first improvement would be a small authenticated API with a database, because every "
            "other production feature depends on a real user and appointment store.",
        ),
        (
            "How do you measure success of this project academically?",
            "Traceability: each specific objective maps to a running feature (Table 7). Success is "
            "not clinical outcome statistics; it is a complete, localised, defensible prototype "
            "with honest limits.",
        ),
    ]
    for i, (q, a) in enumerate(qa, 1):
        r.h(f"C.{i}  {q}", 3)
        r.p(a)

    # ------------------------------------------------------------------
    # APPENDIX D
    # ------------------------------------------------------------------
    r.h("APPENDIX D  FILL-IN CHECKLIST BEFORE PRINTING")
    r.p("Replace every item in square brackets on the cover and declaration pages:", first_line=False)
    r.bullet("University / faculty / department names exactly as the institution writes them.")
    r.bullet("Official degree title (for example Bachelor of Science in Computer Science).")
    r.bullet("Student full names as on the registration card, and registration number.")
    r.bullet("Supervisor names, academic rank and department.")
    r.bullet("Add a university logo on the cover if the department requires it.")
    r.bullet("Update the reference list to the department’s citation style; remove unused placeholders.")
    r.bullet("If the department requires a plagiarism report, generate it from this document plus the source code.")
    r.bullet("Print the abstract on its own page if the faculty template demands a standalone abstract.")
    r.p(
        "This report is written to convince a supervisor that the project is (1) locally relevant, "
        "(2) complete enough to demonstrate, (3) technically explained, (4) ethically cautious, "
        "and (5) honest about what remains for production. Those five points are the argument of the defence.",
        first_line=False,
    )

    r.cover_line("— End of report —", size=12, italic=True, color=TEAL, space_after=4)

    return r


def main():
    out_dir = Path(r"c:\Users\ishim\OneDrive\Desktop\pictures\healthline")
    desktop = Path(r"c:\Users\ishim\OneDrive\Desktop")
    name = "Ubuzima_Bwiza_Project_Defense_Report.docx"
    r = build()
    p1 = out_dir / name
    p2 = desktop / name
    r.save(p1)
    r.save(p2)
    print(f"Wrote {p1}")
    print(f"Wrote {p2}")


if __name__ == "__main__":
    main()
